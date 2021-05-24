# This is the sample program from the python-docx documentation.
# I plan on copying the code, commenting throughout, and iterating to make it my own.

from docx import Document
from docx.shared import Inches
from PIL import Image
import os

os.chdir('/home/pi/mu_code/docxTest')
document = Document()

#This line sets the heading for the document. 0 = largest heading?
document.add_heading('Document Title', 0)

# add_paragraph() means exactly that. add_run() writes a string in some different style.
p = document.add_paragraph('A plain paragraph that\'s got some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italics.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph('first item in unordered list', style='List Bullet')
document.add_paragraph('first items in ordered list', style='List Number')

# Adding a picture. Looks like .add_picture() requires the picture's absolute path to load.
document.add_picture('/home/pi/mu_code/docxTest/the-truth.jpg', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demoDoc.docx')